# -*- coding: utf-8 -*-
"""
StudyForge Blackbook Generator - Complete 90+ Page Academic Report
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

OUTPUT_PATH = r"e:\notex\BLACK_BOOK_MATERIALS\StudyForge_Blackbook_COMPLETE.docx"

doc = Document()

# ============================================================
# STYLES SETUP
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_chapter_heading(text, centered=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_subheading(text, level=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_sub_subheading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body_bold(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5
    return p

def add_code_block(code_text, caption=""):
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.italic = True
        p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code_text)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="4" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="4" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="4" w:color="999999"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def add_figure_placeholder(caption, fig_num):
    # Add bordered box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n[Screenshot / Diagram Placeholder]\n\n")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True
    # Set cell border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)
    # Caption below
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure {fig_num}: {caption}")
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.italic = True
    cap.paragraph_format.space_after = Pt(12)
    return cap

def add_table_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_styled_table(headers, rows, caption=""):
    if caption:
        add_table_caption(caption)
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.bold = True
        set_cell_shading(cell, "D9E2F3")
    # Data rows
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.cell(r_i+1, c_i)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    doc.add_paragraph()  # spacing

def add_page_break():
    doc.add_page_break()

# ============================================================
# FRONT MATTER
# ============================================================

# --- Title Page ---
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("StudyForge: An AI-Powered Document-to-Study-System Pipeline")
run.font.size = Pt(18)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted In Partial Fulfillment of Requirements For the Degree Of")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Masters of Science (Computer Science)")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("By")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Varun R. Darji")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Roll No: 310315240004")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Guide")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prof. Swati Maurya")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.bold = True

for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Somaiya School of Basic and Applied Science\nSomaiya Vidyavihar University\nVidyavihar, Mumbai - 400 077")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2024-26")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.bold = True

add_page_break()

# --- Certificate Page ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Somaiya Vidyavihar University")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Somaiya School of Basic and Applied Science")
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Certificate")
run.font.size = Pt(16)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()
add_body('This is to certify that the project report on "StudyForge: An AI-Powered Document-to-Study-System Pipeline" is a bonafide record of the project work done by Varun R. Darji (Roll No: 310315240004) in the year 2024-26 under the guidance of Prof. Swati Maurya, in partial fulfillment of the requirement for the award of the degree of Masters of Science (Computer Science) from Somaiya Vidyavihar University.')

for _ in range(6):
    doc.add_paragraph()

# Signature table
table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).paragraphs[0].add_run("Guide").font.name = 'Times New Roman'
table.cell(0, 1).paragraphs[0].add_run("Programme Coordinator").font.name = 'Times New Roman'
for _ in range(3):
    doc.add_paragraph()
table2 = doc.add_table(rows=1, cols=2)
table2.cell(0, 0).paragraphs[0].add_run("Head of the Department").font.name = 'Times New Roman'
table2.cell(0, 1).paragraphs[0].add_run("Director").font.name = 'Times New Roman'

doc.add_paragraph()
add_body("Date:")
add_body("Place: Mumbai-77")

add_page_break()

# --- Certificate of Approval ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Somaiya Vidyavihar University")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Somaiya School of Basic and Applied Sciences\nCertificate of Approval of Examiners")
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()
add_body('This is to certify that the project report on "StudyForge: An AI-Powered Document-to-Study-System Pipeline" is a bonafide record of the project work done by Varun R. Darji (Roll No: 310315240004) in partial fulfillment of the requirement for the award of the degree of Masters of Science (Computer Science) from Somaiya Vidyavihar University, and is approved for its content and presentation.')

for _ in range(6):
    doc.add_paragraph()

table3 = doc.add_table(rows=1, cols=2)
table3.cell(0, 0).paragraphs[0].add_run("External Examiner / Expert").font.name = 'Times New Roman'
table3.cell(0, 1).paragraphs[0].add_run("Internal Examiner / Guide").font.name = 'Times New Roman'

doc.add_paragraph()
add_body("Date:")
add_body("Place: Mumbai-77")

add_page_break()

# --- Declaration ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Somaiya Vidyavihar University")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Somaiya School of Basic and Applied Science")
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DECLARATION")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()
add_body('I declare that this written report submission represents the work done based on my and / or others\u2019 ideas with adequately cited and referenced the original source. I also declare that I have adhered to all principles of academic honesty and integrity and have not misrepresented or fabricated or falsified any idea / data / fact / source in my submission.')

add_body('I understand that any violation of the above will be cause for disciplinary action by the college and may evoke penal action from the sources which have not been properly cited or from whom proper permission has not been taken when needed.')

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Signature of the Student: Varun R. Darji")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

p = doc.add_paragraph()
run = p.add_run("Roll No: 310315240004")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()
add_body("Date:")
add_body("Place: Mumbai-77")

add_page_break()

# --- Acknowledgement ---
add_chapter_heading("Acknowledgement")
doc.add_paragraph()
add_body("This project would not have been possible without the valuable guidance and support of many people to whom I am deeply indebted.")
add_body("I express my sincere gratitude to Prof. Swati Maurya, Department of Computer Science, Somaiya School of Basic and Applied Science, for their invaluable guidance, constant encouragement, and constructive feedback throughout the development of this project. Their expertise in artificial intelligence and software engineering shaped the direction of this work significantly.")
add_body("I am thankful to Dr. [Head of Department Name], Head of the Department of Computer Science, for providing the necessary infrastructure and academic environment that facilitated this research.")
add_body("I extend my appreciation to the Director of Somaiya School of Basic and Applied Science for granting the opportunity to undertake this project as part of the Master of Science programme.")
add_body("I also wish to acknowledge the open-source community, particularly the developers of the Notex project, Go programming language ecosystem, and Google\u2019s Gemini API, whose tools and frameworks formed the technological foundation of StudyForge.")
add_body("Finally, I thank my family and friends for their unwavering support and patience throughout the duration of this project.")

for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Varun R. Darji")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

add_page_break()

# --- Abstract ---
add_chapter_heading("Abstract")
doc.add_paragraph()
add_body("Traditional note-taking applications function primarily as passive storage systems, offering basic organisational features like folders and tags but lacking the intelligence to transform raw academic content into structured, study-optimised material. While recent AI-powered tools such as Google\u2019s NotebookLM and Notion AI have introduced document-level chat and summarisation, they operate as shallow, single-layer systems that do not address the deeper challenges of knowledge synthesis, multi-format note generation, and hierarchical content organisation.")

add_body("This project presents StudyForge, an AI-powered document-to-study-system pipeline that extends the capabilities of existing tools through a multi-layered knowledge processing architecture. Built upon the open-source Notex framework, StudyForge introduces four significant enhancements: (1) migration from OpenAI to Google\u2019s Gemini 1.5 Flash language model via an OpenAI-compatible API endpoint, enabling cost-effective and high-performance LLM integration; (2) a multi-mode note generation engine supporting four distinct transformation types \u2014 Summary, FAQ, Study Guide, and Exam Notes \u2014 each designed for a specific study context and implemented via targeted prompt engineering; (3) a Notebook organisational system for grouping related sources into named, colour-coded knowledge containers; and (4) a Textbook Generation module that compiles all sources and notes within a notebook into a chapter-structured study document.")

add_body("The system is implemented as a monolithic Go backend serving a vanilla JavaScript single-page application frontend, with SQLite as the persistent data store. Document ingestion supports PDF, DOCX, TXT, Markdown, HTML, and image-based files through a combination of the Markitdown library and Gemini Vision OCR. A Retrieval-Augmented Generation (RAG) chat system enables users to ask natural language questions and receive cited, source-grounded answers drawn from their entire knowledge base.")

add_body("Functional testing across all modules validated ingestion accuracy, note generation quality, retrieval precision, and chat response grounding. The system processed documents ranging from 1 to 50 pages with note generation completing within 30 seconds and search query latency remaining under 2 seconds for all tested corpus sizes.")

add_body("Keywords: Retrieval-Augmented Generation, Large Language Models, Knowledge Synthesis, Educational Technology, Document Processing, Gemini API, Go, Single-Page Application")

add_page_break()

# --- Table of Contents ---
add_chapter_heading("Table of Contents")
doc.add_paragraph()
toc_items = [
    ("", "Certificate", "i"),
    ("", "Certificate of Approval of Examiners", "ii"),
    ("", "Declaration", "iii"),
    ("", "Acknowledgement", "iv"),
    ("", "Abstract", "v"),
    ("", "Table of Contents", "vi"),
    ("", "List of Figures", "viii"),
    ("", "List of Tables", "ix"),
    ("1", "Introduction", "1"),
    ("1.1", "Introduction", "1"),
    ("1.2", "Background", "2"),
    ("1.3", "Objectives", "4"),
    ("1.4", "Purpose, Scope and Applicability", "5"),
    ("1.5", "Organisation of the Report", "7"),
    ("2", "Literature Survey", "8"),
    ("2.1", "Overview of Related Work", "8"),
    ("2.2", "Review of Existing Technologies", "9"),
    ("2.3", "Identified Research Gaps", "15"),
    ("2.4", "Objectives of Thesis Work", "17"),
    ("3", "Requirements and Analysis", "18"),
    ("3.1", "Problem Definition", "18"),
    ("3.2", "Requirements Specification", "19"),
    ("3.3", "Project SDLC Model", "22"),
    ("3.4", "Planning and Scheduling", "23"),
    ("3.5", "Software and Hardware Requirements", "24"),
    ("3.6", "Technology Stack", "25"),
    ("3.7", "Conceptual Models", "29"),
    ("4", "System Design", "32"),
    ("4.1", "System Architecture", "32"),
    ("4.2", "Database Design", "34"),
    ("4.3", "Procedural Design", "38"),
    ("4.4", "User Interface Design", "41"),
    ("4.5", "API Design", "44"),
    ("5", "Implementation and Testing", "47"),
    ("5.1", "Implementation Approach", "47"),
    ("5.2", "Coding Details", "49"),
    ("5.3", "Frontend Implementation", "56"),
    ("5.4", "Testing", "59"),
    ("6", "Results and Discussion", "64"),
    ("6.1", "Module-Wise Results", "64"),
    ("6.2", "Performance Evaluation", "68"),
    ("6.3", "Discussion", "70"),
    ("7", "Conclusions and Future Work", "72"),
    ("7.1", "Conclusions", "72"),
    ("7.2", "Limitations of the System", "73"),
    ("7.3", "Future Scope of the Project", "75"),
    ("7.4", "Semester 2 Roadmap", "77"),
    ("", "References", "79"),
    ("", "Appendix A \u2014 Algorithms and Pseudocode", "82"),
    ("", "Appendix B \u2014 Source Code Snippets", "85"),
    ("", "Appendix C \u2014 Design System Specification", "90"),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    if num and not num.startswith(" "):
        prefix = f"{num}  " if '.' not in num else f"    {num}  "
    else:
        prefix = ""
    is_chapter = bool(num and '.' not in num)
    run = p.add_run(f"{prefix}{title}")
    run.font.size = Pt(12) if is_chapter else Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True if is_chapter else None
    # Right-aligned page number
    run2 = p.add_run(f"\t{page}")
    run2.font.size = Pt(11)
    run2.font.name = 'Times New Roman'

add_page_break()

# --- List of Figures ---
add_chapter_heading("List of Figures")
doc.add_paragraph()
figures = [
    ("3.1", "Use Case Diagram"),
    ("3.2", "Data Flow Diagram \u2014 Level 0 (Context Diagram)"),
    ("3.3", "Data Flow Diagram \u2014 Level 1"),
    ("3.4", "Entity-Relationship Diagram"),
    ("3.5", "Gantt Chart \u2014 Project Timeline"),
    ("4.1", "System Architecture Diagram"),
    ("4.2", "Component Interaction Diagram"),
    ("4.3", "Sequence Diagram \u2014 Document Upload and Processing"),
    ("4.4", "Sequence Diagram \u2014 Note Generation"),
    ("4.5", "Sequence Diagram \u2014 RAG Chat"),
    ("5.1", "Screenshot \u2014 Dashboard View"),
    ("5.2", "Screenshot \u2014 Sources List"),
    ("5.3", "Screenshot \u2014 Source Detail and Note Viewer"),
    ("5.4", "Screenshot \u2014 Notebooks Grid"),
    ("5.5", "Screenshot \u2014 Textbook Viewer"),
    ("5.6", "Screenshot \u2014 Chat Interface"),
    ("5.7", "Screenshot \u2014 Search Results"),
    ("5.8", "Screenshot \u2014 Dark Glassmorphism Theme"),
    ("6.1", "Note Generation Response Times"),
    ("6.2", "Search Latency vs Corpus Size"),
]
for num, title in figures:
    p = doc.add_paragraph()
    run = p.add_run(f"Figure {num}: {title}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

add_page_break()

# --- List of Tables ---
add_chapter_heading("List of Tables")
doc.add_paragraph()
tables_list = [
    ("2.1", "Comparative Analysis of Existing Systems"),
    ("3.1", "Functional Requirements"),
    ("3.2", "Non-Functional Requirements"),
    ("3.3", "Software Requirements"),
    ("3.4", "Hardware Requirements"),
    ("3.5", "Technology Stack Overview"),
    ("4.1", "Sources Table Schema"),
    ("4.2", "Notes Table Schema"),
    ("4.3", "Notebooks Table Schema"),
    ("4.4", "Embeddings / Vector Store Schema"),
    ("4.5", "Chat Sessions and Messages Schema"),
    ("4.6", "Textbooks Table Schema"),
    ("4.7", "API Endpoints \u2014 Sources Module"),
    ("4.8", "API Endpoints \u2014 Notes Module"),
    ("4.9", "API Endpoints \u2014 Notebooks Module"),
    ("4.10", "API Endpoints \u2014 Chat and Search Modules"),
    ("5.1", "Test Cases \u2014 All Modules"),
    ("6.1", "Performance Metrics Summary"),
]
for num, title in tables_list:
    p = doc.add_paragraph()
    run = p.add_run(f"Table {num}: {title}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

add_page_break()

print("[1/5] Front matter complete")

# Save checkpoint
doc.save(OUTPUT_PATH)
print(f"[CHECKPOINT] Saved front matter to {OUTPUT_PATH}")
